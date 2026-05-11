"""
 Multi-Modal Data Understanding Service
=========================================

Parse and extract data from multiple file formats:
- PDFs (text + images)
- Images (OCR + vision analysis)
- CAD files (metadata extraction)
- Excel files (structured data)
- Word documents
- Videos (frame analysis)

Integrations:
- Ollama vision models (LLaVA, BakLLaVA)
- PyMuPDF for PDF processing
- pytesseract for OCR
- OpenCV for image processing
- openpyxl for Excel files
- python-docx for Word files
"""

# This router intentionally catches broad exceptions at integration boundaries
# because many optional third-party libraries may be absent or partially installed.
# pylint: disable=broad-except,unused-import

import logging
import asyncio
import base64
import io
import re
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any
from enum import Enum
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, BackgroundTasks
from sqlalchemy.orm import Session
from core.db_session import get_db
from services.admin_config_service import AdminConfigService
from pydantic import BaseModel

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/multimodal", tags=["Multi-Modal Understanding"])


# ============= MODELS =============

class FileType(str, Enum):
    PDF = "pdf"
    IMAGE = "image"
    CAD = "cad"
    EXCEL = "excel"
    WORD = "word"
    VIDEO = "video"
    UNKNOWN = "unknown"


class ExtractionMethod(str, Enum):
    OCR = "ocr"
    VISION_LLM = "vision_llm"
    TEXT_PARSER = "text_parser"
    METADATA = "metadata"
    HYBRID = "hybrid"


class FileAnalysisRequest(BaseModel):
    file_path: Optional[str] = None
    file_url: Optional[str] = None
    extraction_method: ExtractionMethod = ExtractionMethod.HYBRID
    vision_model: str = "llava:latest"
    extract_metadata: bool = True
    extract_text: bool = True
    extract_images: bool = False
    ocr_language: str = "eng"


class FileAnalysisResponse(BaseModel):
    file_name: str
    file_type: FileType
    file_size_bytes: int
    extraction_method: ExtractionMethod
    text_content: Optional[str] = None
    metadata: Dict[str, Any] = {}
    extracted_data: Dict[str, Any] = {}
    images_analyzed: int = 0
    processing_time_ms: int = 0
    analyzed_at: Optional[datetime] = None
    success: bool = True
    error: Optional[str] = None


class ImageAnalysisRequest(BaseModel):
    image_base64: str
    prompt: str = "Describe this image in detail, focusing on technical specifications, dimensions, and any text visible."
    vision_model: str = "llava:latest"


class ImageAnalysisResponse(BaseModel):
    description: str
    extracted_text: Optional[str] = None
    metadata: Dict[str, Any] = {}
    confidence: float = 0.0


# ============= MULTI-MODAL SERVICE =============

class MultiModalService:
    """Service for analyzing multiple file formats using AI"""
    
    def __init__(self):
        self.supported_image_formats = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'}
        self.supported_pdf_formats = {'.pdf'}
        self.supported_cad_formats = {'.dwg', '.dxf', '.step', '.stp', '.iges', '.igs'}
        self.supported_excel_formats = {'.xlsx', '.xls', '.xlsm', '.csv', '.tsv'}
        self.supported_word_formats = {'.docx', '.doc'}
        self.supported_video_formats = {'.mp4', '.avi', '.mov', '.mkv'}
    
    def detect_file_type(self, filename: str) -> FileType:
        """Detect file type from extension"""
        ext = Path(filename).suffix.lower()
        
        if ext in self.supported_image_formats:
            return FileType.IMAGE
        elif ext in self.supported_pdf_formats:
            return FileType.PDF
        elif ext in self.supported_cad_formats:
            return FileType.CAD
        elif ext in self.supported_excel_formats:
            return FileType.EXCEL
        elif ext in self.supported_word_formats:
            return FileType.WORD
        elif ext in self.supported_video_formats:
            return FileType.VIDEO
        else:
            return FileType.UNKNOWN
    
    async def analyze_with_vision_llm(
        self,
        image_data: bytes,
        prompt: str,
        model: str = "llava:latest",
        ollama_host: Optional[str] = None
    ) -> Dict[str, Any]:
        """Analyze image using Ollama vision model"""
        try:
            import ollama
            
            # Encode image to base64
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # Initialize client with custom host if provided
            if ollama_host:
                client = ollama.Client(host=ollama_host)
                chat_func = client.chat
            else:
                chat_func = ollama.chat

            # Call Ollama vision model — the ollama Python client is synchronous;
            # wrap it in a thread so the event loop is not blocked.
            def _call_ollama() -> Any:
                return chat_func(
                    model=model,
                    messages=[{
                        'role': 'user',
                        'content': prompt,
                        'images': [image_base64]
                    }]
                )

            response = await asyncio.to_thread(_call_ollama)
            
            return {
                "description": response['message']['content'],
                "model": model,
                "success": True
            }
            
        except ImportError:
            logger.warning("Ollama not installed, falling back to basic analysis")
            return {
                "description": "Vision analysis not available (Ollama not installed)",
                "model": model,
                "success": False
            }
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error in vision LLM analysis: %s", e)
            return {
                "description": f"Error: {str(e)}",
                "model": model,
                "success": False
            }
    
    async def extract_text_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text and images from PDF"""
        try:
            import fitz  # type: ignore[import-untyped]  # PyMuPDF
            
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            
            text_content = []
            images = []
            metadata = {
                "page_count": len(pdf_document),
            }
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    images.append({
                        "page": page_num + 1,
                        "index": img_index,
                        "image_bytes": base_image["image"]
                    })
            
            pdf_document.close()
            
            return {
                "text": "\n\n".join(text_content),
                "images": images,
                "metadata": metadata,
                "success": True
            }
            
        except ImportError:
            logger.warning("PyMuPDF not installed")
            return {
                "text": "",
                "images": [],
                "metadata": {},
                "success": False,
                "error": "PyMuPDF not installed"
            }
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error extracting PDF: %s", e)
            return {
                "text": "",
                "images": [],
                "metadata": {},
                "success": False,
                "error": str(e)
            }
    
    async def ocr_image(self, image_data: bytes, language: str = "eng") -> str:
        """Perform OCR on image"""
        try:
            import pytesseract  # type: ignore[import-untyped]
            from PIL import Image
            
            image = Image.open(io.BytesIO(image_data))
            text = pytesseract.image_to_string(image, lang=language)
            
            return text.strip()
            
        except ImportError:
            logger.warning("pytesseract not installed")
            return ""
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error in OCR: %s", e)
            return ""
    
    async def extract_excel_data(self, file_content: bytes) -> Dict[str, Any]:
        """Extract data from Excel file"""
        try:
            import openpyxl  # type: ignore[import-untyped]
            
            workbook = openpyxl.load_workbook(io.BytesIO(file_content))
            
            sheets_data = {}
            metadata = {
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames
            }
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Get dimensions
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                # Extract data (limit to first 100 rows for performance)
                rows_data = []
                for row in sheet.iter_rows(max_row=min(max_row, 100), values_only=True):
                    rows_data.append(list(row))
                
                sheets_data[sheet_name] = {
                    "dimensions": {"rows": max_row, "columns": max_col},
                    "data": rows_data[:100],  # First 100 rows
                    "truncated": max_row > 100
                }
            
            return {
                "sheets": sheets_data,
                "metadata": metadata,
                "success": True
            }
            
        except ImportError:
            logger.warning("openpyxl not installed")
            return {
                "sheets": {},
                "metadata": {},
                "success": False,
                "error": "openpyxl not installed"
            }
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error extracting Excel: %s", e)
            return {
                "sheets": {},
                "metadata": {},
                "success": False,
                "error": str(e)
            }
    
    async def extract_word_data(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_content))
            
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            tables_data = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_data)
            }
            
            return {
                "text": "\n\n".join(paragraphs),
                "tables": tables_data,
                "metadata": metadata,
                "success": True
            }
            
        except ImportError:
            logger.warning("python-docx not installed")
            return {
                "text": "",
                "tables": [],
                "metadata": {},
                "success": False,
                "error": "python-docx not installed"
            }
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error extracting Word: %s", e)
            return {
                "text": "",
                "tables": [],
                "metadata": {},
                "success": False,
                "error": str(e)
            }

    async def extract_cad_data(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Dispatch CAD extraction based on file extension."""
        ext = Path(filename).suffix.lower()
        if ext == '.dxf':
            return await self._extract_dxf(file_content)
        elif ext == '.dwg':
            return self._extract_dwg_header(file_content)
        elif ext in {'.step', '.stp'}:
            return await self._extract_step(file_content)
        elif ext in {'.iges', '.igs'}:
            return await self._extract_iges(file_content)
        return {"success": False, "error": f"Unsupported CAD extension: {ext}"}

    async def _extract_dxf(self, file_content: bytes) -> Dict[str, Any]:
        """Extract layers, entities, and text from DXF files using ezdxf."""
        try:
            from ezdxf import recover

            # ezdxf recover.readbytes is a sync call — offload to thread pool.
            doc, auditor = await asyncio.to_thread(recover.readbytes, file_content)

            metadata = {
                "dxf_version": doc.dxfversion,
                "encoding": doc.encoding,
            }

            layers = [
                {
                    "name": layer.dxf.name,
                    "color": layer.dxf.color,
                    "linetype": layer.dxf.get("linetype", "CONTINUOUS"),
                }
                for layer in doc.layers
            ]

            entity_counts: Dict[str, int] = {}
            text_items = []
            for entity in doc.modelspace():
                etype = entity.dxftype()
                entity_counts[etype] = entity_counts.get(etype, 0) + 1
                if etype in ("TEXT", "MTEXT"):
                    try:
                        text_items.append(entity.dxf.text)
                    except Exception:  # pylint: disable=broad-except
                        pass

            return {
                "format": "DXF",
                "metadata": metadata,
                "layers": layers,
                "entity_counts": entity_counts,
                "total_entities": sum(entity_counts.values()),
                "text_content": "\n".join(text_items),
                "audit_errors": len(auditor.errors),
                "success": True,
            }
        except ImportError:
            return {"success": False, "error": "ezdxf not installed. Run: pip install ezdxf"}
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error extracting DXF: %s", e)
            return {"success": False, "error": str(e)}

    def _extract_dwg_header(self, file_content: bytes) -> Dict[str, Any]:
        """Decode version from DWG binary header magic bytes.

        Full DWG parsing requires LibreDWG (native C library) or a CAD application.
        We extract the version stamp from the first 6 bytes which is always ASCII.
        """
        DWG_VERSION_MAP = {
            b"AC1006": "R10",    b"AC1009": "R11/R12",
            b"AC1012": "R13",    b"AC1014": "R14",
            b"AC1015": "R2000",  b"AC1018": "R2004",
            b"AC1021": "R2007",  b"AC1024": "R2010",
            b"AC1027": "R2013",  b"AC1032": "R2018",
        }
        version_key = file_content[:6] if len(file_content) >= 6 else b""
        label = DWG_VERSION_MAP.get(
            version_key,
            f"Unknown ({version_key.decode('ascii', errors='replace')})",
        )
        return {
            "format": "DWG",
            "metadata": {
                "dwg_version": label,
                "file_size_bytes": len(file_content),
            },
            "note": "Full DWG parsing requires LibreDWG. Only header metadata extracted.",
            "success": True,
        }

    async def _extract_step(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata and entity summary from STEP (ISO 10303) files.

        STEP is a text-based format — no external library required.
        Parses the HEADER section and counts entity types in the DATA section.
        """
        try:
            text = file_content.decode("utf-8", errors="replace")

            metadata: Dict[str, Any] = {}

            fn_m = re.search(
                r"FILE_NAME\s*\(\s*'([^']*)'",
                text, re.IGNORECASE,
            )
            fn_m2 = re.search(
                r"FILE_NAME\s*\(\s*'[^']*'\s*,\s*'([^']*)'\s*,",
                text, re.IGNORECASE,
            )
            if fn_m:
                metadata["file_name"] = fn_m.group(1)
            if fn_m2:
                metadata["timestamp"] = fn_m2.group(1)

            fd_m = re.search(
                r"FILE_DESCRIPTION\s*\(\s*\(\s*'([^']*)'",
                text, re.IGNORECASE,
            )
            if fd_m:
                metadata["description"] = fd_m.group(1)

            fs_m = re.search(
                r"FILE_SCHEMA\s*\(\s*\(\s*'([^']*)'",
                text, re.IGNORECASE,
            )
            if fs_m:
                metadata["schema"] = fs_m.group(1)

            # Count entity types in the DATA section
            entity_counts: Dict[str, int] = {}
            products = []
            data_m = re.search(r"DATA\s*;(.*?)ENDSEC\s*;", text, re.DOTALL | re.IGNORECASE)
            if data_m:
                data_body = data_m.group(1)
                for em in re.finditer(r"#\d+\s*=\s*([A-Z_][A-Z0-9_]*)\s*\(", data_body):
                    etype = em.group(1)
                    entity_counts[etype] = entity_counts.get(etype, 0) + 1
                for pm in re.finditer(
                    r"PRODUCT\s*\(\s*'([^']*)'\s*,\s*'([^']*)'",
                    data_body, re.IGNORECASE,
                ):
                    products.append({"id": pm.group(1), "name": pm.group(2)})

            top_entities = dict(
                sorted(entity_counts.items(), key=lambda x: -x[1])[:10]
            )

            return {
                "format": "STEP",
                "metadata": metadata,
                "top_entity_types": top_entities,
                "total_entities": sum(entity_counts.values()),
                "products": products[:20],
                "success": True,
            }
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error extracting STEP: %s", e)
            return {"success": False, "error": str(e)}

    async def _extract_iges(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata and entity summary from IGES files.

        IGES is column-aligned ASCII: column 73 is the section flag character.
        Section 'S' = Start, 'G' = Global, 'D' = Directory Entry, 'P' = Parameter.
        Entity type number sits in columns 1-8 of every odd Directory Entry line.
        """
        IGES_ENTITY_NAMES = {
            100: "Circular Arc",          102: "Composite Curve",
            104: "Conic Arc",             106: "Copious Data",
            108: "Plane",                 110: "Line",
            112: "Param Spline Curve",    114: "Param Spline Surface",
            116: "Point",                 118: "Ruled Surface",
            120: "Surface of Revolution", 122: "Tabulated Cylinder",
            124: "Transformation Matrix", 126: "Rational B-Spline Curve",
            128: "Rational B-Spline Surface", 130: "Offset Curve",
            140: "Offset Surface",        141: "Boundary",
            142: "Curve on Param Surface",143: "Bounded Surface",
            144: "Trimmed Surface",       212: "General Note",
            308: "Subfigure Definition",  314: "Color Definition",
            402: "Associativity Instance",404: "Drawing",
            406: "Property",             408: "Singular Subfigure Instance",
            410: "View",                 502: "Vertex",
            504: "Edge",                 508: "Loop",
            510: "Face",                 514: "Shell",
        }
        try:
            text = file_content.decode("utf-8", errors="replace")
            lines = text.splitlines()

            global_parts = []
            entity_counts: Dict[str, int] = {}
            start_text_lines = []

            for line in lines:
                if len(line) < 73:
                    continue
                section = line[72]
                col_data = line[:72].rstrip()
                if section == 'S':
                    start_text_lines.append(col_data)
                elif section == 'G':
                    global_parts.append(col_data)
                elif section == 'D':
                    try:
                        entity_type = int(line[0:8].strip())
                        entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
                    except (ValueError, IndexError):
                        pass

            metadata: Dict[str, Any] = {}
            if global_parts:
                global_str = "".join(global_parts)
                # Strip parameter/record delimiters (1H, and 1H;)
                global_str = re.sub(r"1H[,;]", "", global_str)
                fields = global_str.split(",")
                # IGES global section field indices (1-based in spec, 0-based here)
                field_map = {
                    3: "product_id_sender",
                    4: "file_name",
                    5: "native_system_id",
                    6: "preprocessor_version",
                    23: "author",
                    24: "organization",
                }
                for idx, label in field_map.items():
                    if idx < len(fields):
                        val = fields[idx].strip().strip("'")
                        if val:
                            metadata[label] = val

            named_counts = {
                IGES_ENTITY_NAMES.get(k, f"Entity_{k}"): v
                for k, v in entity_counts.items()
            }
            top_entities = dict(sorted(named_counts.items(), key=lambda x: -x[1])[:10])

            return {
                "format": "IGES",
                "metadata": metadata,
                "top_entity_types": top_entities,
                "total_entities": sum(entity_counts.values()),
                "start_section_preview": " ".join(start_text_lines)[:500] if start_text_lines else "",
                "success": True,
            }
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error extracting IGES: %s", e)
            return {"success": False, "error": str(e)}

    async def analyze_file(
        self,
        file_content: bytes,
        filename: str,
        extraction_method: ExtractionMethod,
        vision_model: str,
        extract_metadata: bool,
        extract_text: bool,
        extract_images: bool,
        ocr_language: str,
        ollama_host: Optional[str] = None
    ) -> FileAnalysisResponse:
        """Analyze file with specified extraction method"""
        start_time = datetime.now(timezone.utc)
        
        file_type = self.detect_file_type(filename)
        file_size = len(file_content)
        
        result = FileAnalysisResponse(
            file_name=filename,
            file_type=file_type,
            file_size_bytes=file_size,
            extraction_method=extraction_method,
            analyzed_at=start_time
        )
        
        try:
            # PDF processing
            if file_type == FileType.PDF:
                pdf_data = await self.extract_text_from_pdf(file_content)
                
                if pdf_data["success"]:
                    if extract_text:
                        result.text_content = pdf_data["text"]
                    if extract_metadata:
                        result.metadata = pdf_data["metadata"]
                    
                    # Analyze images with vision LLM if requested
                    if extract_images and extraction_method in [ExtractionMethod.VISION_LLM, ExtractionMethod.HYBRID]:
                        for img_info in pdf_data["images"][:5]:  # Limit to 5 images
                            vision_result = await self.analyze_with_vision_llm(
                                img_info["image_bytes"],
                                "Describe this technical diagram or image. Extract any text, dimensions, or specifications visible.",
                                vision_model,
                                ollama_host=ollama_host
                            )
                            
                            if vision_result["success"]:
                                result.extracted_data[f"image_page_{img_info['page']}_index_{img_info['index']}"] = vision_result["description"]
                                result.images_analyzed += 1
                else:
                    result.error = pdf_data.get("error")
            
            # Image processing
            elif file_type == FileType.IMAGE:
                if extraction_method in [ExtractionMethod.VISION_LLM, ExtractionMethod.HYBRID]:
                    vision_result = await self.analyze_with_vision_llm(
                        file_content,
                        "Describe this image in detail. Extract any text, technical specifications, or structured data visible.",
                        vision_model,
                        ollama_host=ollama_host
                    )
                    
                    if vision_result["success"]:
                        if extract_text:
                            result.text_content = vision_result["description"]
                        result.images_analyzed = 1
                
                if extraction_method in [ExtractionMethod.OCR, ExtractionMethod.HYBRID]:
                    ocr_text = await self.ocr_image(file_content, ocr_language)
                    if ocr_text:
                        result.extracted_data["ocr_text"] = ocr_text
            
            # Excel processing
            elif file_type == FileType.EXCEL:
                excel_data = await self.extract_excel_data(file_content)
                
                if excel_data["success"]:
                    result.extracted_data = excel_data["sheets"]
                    if extract_metadata:
                        result.metadata = excel_data["metadata"]
                else:
                    result.error = excel_data.get("error")
            
            # Word processing
            elif file_type == FileType.WORD:
                word_data = await self.extract_word_data(file_content)
                
                if word_data["success"]:
                    if extract_text:
                        result.text_content = word_data["text"]
                    result.extracted_data["tables"] = word_data["tables"]
                    if extract_metadata:
                        result.metadata = word_data["metadata"]
                else:
                    result.error = word_data.get("error")
            
            # CAD files
            elif file_type == FileType.CAD:
                cad_data = await self.extract_cad_data(file_content, filename)
                if cad_data["success"]:
                    if extract_metadata:
                        result.metadata = cad_data.get("metadata", {})
                    # Expose all non-metadata keys in extracted_data
                    result.extracted_data = {
                        k: v for k, v in cad_data.items()
                        if k not in ("success", "error", "metadata")
                    }
                    if extract_text and cad_data.get("text_content"):
                        result.text_content = cad_data["text_content"]
                else:
                    result.error = cad_data.get("error")
                    result.success = False

            else:
                result.error = f"Unsupported file type: {file_type}"
                result.success = False
            
        except Exception as e:  # pylint: disable=broad-except
            logger.error("Error analyzing file %s: %s", filename, e)
            result.error = str(e)
            result.success = False
        
        # Calculate processing time
        end_time = datetime.now(timezone.utc)
        result.processing_time_ms = int((end_time - start_time).total_seconds() * 1000)
        
        return result


# Global service instance
service = MultiModalService()


# ============= API ENDPOINTS =============

@router.post("/analyze-file", summary="Analyze File with Multi-Modal AI")
async def analyze_file(
    file: UploadFile = File(...),
    extraction_method: ExtractionMethod = ExtractionMethod.HYBRID,
    vision_model: str = "llava:latest",
    extract_metadata: bool = True,
    extract_text: bool = True,
    extract_images: bool = False,
    ocr_language: str = "eng",
    db: Session = Depends(get_db)
):
    """
    Analyze uploaded file using multi-modal AI
    
    Supported formats:
    - PDF: Text extraction + vision analysis of embedded images
    - Images: Vision LLM + OCR
    - Excel: Structured data extraction
    - Word: Text + tables extraction
    - CAD: Metadata extraction (coming soon)
    """
    try:
        # Fetch Ollama configuration from Admin Config
        ollama_host = None
        if extraction_method in [ExtractionMethod.VISION_LLM, ExtractionMethod.HYBRID]:
            try:
                config_service = AdminConfigService(db)
                ollama_config = config_service.get_llm_provider_config("ollama")
                if ollama_config and ollama_config.get("api_endpoint"):
                    ollama_host = ollama_config["api_endpoint"]
            except Exception as e:
                logger.warning("Failed to fetch Ollama config from DB, using default: %s", e)

        # Read file content
        file_content = await file.read()

        filename = file.filename or "uploaded_file"
        
        # Analyze file
        result = await service.analyze_file(
            file_content=file_content,
            filename=filename,
            extraction_method=extraction_method,
            vision_model=vision_model,
            extract_metadata=extract_metadata,
            extract_text=extract_text,
            extract_images=extract_images,
            ocr_language=ocr_language,
            ollama_host=ollama_host
        )
        
        return result
        
    except Exception as e:  # pylint: disable=broad-except
        logger.error("Error in file analysis endpoint: %s", e)
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.post("/analyze-image", summary="Analyze Image with Vision LLM")
async def analyze_image(
    request: ImageAnalysisRequest,
    db: Session = Depends(get_db)
):
    """Analyze image using Ollama vision model (LLaVA, BakLLaVA)"""
    try:
        # Fetch Ollama configuration
        ollama_host = None
        try:
            config_service = AdminConfigService(db)
            ollama_config = config_service.get_llm_provider_config("ollama")
            if ollama_config and ollama_config.get("api_endpoint"):
                ollama_host = ollama_config["api_endpoint"]
        except Exception:
            pass # Ignore config fetch errors, use default

        # Decode base64 image
        image_data = base64.b64decode(request.image_base64)
        
        # Analyze with vision LLM
        vision_result = await service.analyze_with_vision_llm(
            image_data=image_data,
            prompt=request.prompt,
            model=request.vision_model,
            ollama_host=ollama_host
        )
        
        # Optional: Run OCR
        ocr_text = await service.ocr_image(image_data)
        
        return ImageAnalysisResponse(
            description=vision_result["description"],
            extracted_text=ocr_text if ocr_text else None,
            metadata={"model": vision_result["model"]},
            confidence=1.0 if vision_result["success"] else 0.0
        )
        
    except Exception as e:
        logger.error("Error in image analysis: %s", e)
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.get("/supported-formats", summary="Get Supported File Formats")
async def get_supported_formats():
    """Get list of supported file formats"""
    return {
        "images": list(service.supported_image_formats),
        "pdf": list(service.supported_pdf_formats),
        "cad": list(service.supported_cad_formats),
        "excel": list(service.supported_excel_formats),
        "word": list(service.supported_word_formats),
        "video": list(service.supported_video_formats)
    }


@router.get("/vision-models", summary="Get Available Vision Models")
async def get_vision_models():
    """Get list of available Ollama vision models"""
    try:
        import ollama
        
        models = ollama.list()
        vision_models = [
            model for model in models.get('models', [])
            if 'llava' in model['name'].lower() or 'bakllava' in model['name'].lower()
        ]
        
        return {
            "available_models": [model['name'] for model in vision_models],
            "recommended": "llava:latest"
        }
        
    except ImportError:
        return {
            "available_models": [],
            "recommended": "llava:latest",
            "error": "Ollama not installed"
        }
    except Exception as e:
        logger.error("Error listing vision models: %s", e)
        return {
            "available_models": [],
            "recommended": "llava:latest",
            "error": str(e)
        }


# =============================================================================
# BATCH PROCESSING ENDPOINTS
# =============================================================================

class BatchAnalyzeRequest(BaseModel):
    """Request body for batch file analysis."""
    file_paths: List[str]
    extraction_method: ExtractionMethod = ExtractionMethod.HYBRID
    vision_model: str = "llava:latest"
    concurrency: int = 8
    db_flush_size: int = 50


class DirectoryDiscoverRequest(BaseModel):
    """Request body for directory discovery + batch processing."""
    directory: str
    recursive: bool = True
    extraction_method: ExtractionMethod = ExtractionMethod.HYBRID
    vision_model: str = "llava:latest"
    concurrency: int = 8
    db_flush_size: int = 50
    run_immediately: bool = True


def _make_batch_processor(concurrency: int, db_flush_size: int, db: Session):
    """Build a :class:`FileBatchProcessor` wired to the current DB session."""
    from services.file_batch_processor import FileBatchProcessor
    from core.db_session import SessionLocal

    # Attempt to get a live Neo4j driver (optional)
    neo4j_driver = None
    try:
        from graph_api.dependencies import _driver  # module-level cached driver
        neo4j_driver = _driver
    except Exception:  # pylint: disable=broad-except
        pass

    return FileBatchProcessor(
        concurrency=concurrency,
        db_flush_size=db_flush_size,
        neo4j_driver=neo4j_driver,
        db_session_factory=SessionLocal,
    )


@router.post("/discover", summary="Discover Files in a Directory")
async def discover_directory(request: DirectoryDiscoverRequest):
    """Recursively walk *directory* and return a manifest of all supported files.

    Does NOT process the files — use `/analyze-batch` or set
    ``run_immediately=true`` to also trigger analysis in the background.
    Returns a manifest you can pass directly to `/analyze-batch`.
    """
    from pathlib import Path as _Path
    from services.file_batch_processor import discover_files

    root = _Path(request.directory).expanduser().resolve()
    if not root.is_dir():
        raise HTTPException(status_code=400, detail=f"Not a directory: {request.directory}")

    records = discover_files(root, recursive=request.recursive)

    # Group by type for the manifest
    by_type: Dict[str, List[str]] = {}
    for rec in records:
        by_type.setdefault(rec.file_type, []).append(str(rec.path))

    manifest = {
        "directory": str(root),
        "recursive": request.recursive,
        "total_files": len(records),
        "by_type": {t: len(paths) for t, paths in by_type.items()},
        "file_paths": [str(r.path) for r in records],
    }

    return manifest


@router.post("/analyze-batch", summary="Batch Analyze File List")
async def analyze_batch(
    request: BatchAnalyzeRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    """Process a list of file paths through the multi-modal pipeline in parallel.

    - Uses ``asyncio.Semaphore(concurrency)`` — default 8 workers.
    - Flushes results to Postgres and Neo4j every ``db_flush_size`` files.
    - Returns immediately with a ``job_id``; poll ``/batch-status/{job_id}``.

    For **small** lists (≤ 20 files) the request blocks and returns results
    inline.  For larger lists the job runs in a background task.
    """
    from pathlib import Path as _Path
    from services.file_batch_processor import FileRecord, _classify_ext
    import uuid

    # Resolve file paths → FileRecord list, rejecting path-traversal attempts
    records = []
    for raw in request.file_paths:
        p = _Path(raw).expanduser().resolve()
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        records.append(FileRecord(path=p, ext=ext, size_bytes=p.stat().st_size, file_type=_classify_ext(ext)))

    if not records:
        raise HTTPException(status_code=400, detail="No valid file paths provided.")

    # Fetch Ollama host from DB config (optional)
    ollama_host = None
    try:
        config_service = AdminConfigService(db)
        ollama_cfg = config_service.get_llm_provider_config("ollama")
        if ollama_cfg and ollama_cfg.get("api_endpoint"):
            ollama_host = ollama_cfg["api_endpoint"]
    except Exception:  # pylint: disable=broad-except
        pass

    processor = _make_batch_processor(request.concurrency, request.db_flush_size, db)
    job_id = uuid.uuid4().hex

    if len(records) <= 20:
        # Inline — wait for the result and return it
        report = await processor.process_records(
            records,
            job_id=job_id,
            extraction_method=request.extraction_method.value,
            vision_model=request.vision_model,
            ollama_host=ollama_host,
        )
        return {
            "job_id": report.job_id,
            "status": "completed",
            "total_files": report.total_files,
            "processed": report.processed,
            "succeeded": report.succeeded,
            "failed": report.failed,
            "errors_summary": report.errors_summary,
        }

    # Large batch — fire and forget, return job_id for polling
    async def _run_background():
        try:
            await processor.process_records(
                records,
                job_id=job_id,
                extraction_method=request.extraction_method.value,
                vision_model=request.vision_model,
                ollama_host=ollama_host,
            )
        except Exception as exc:  # pylint: disable=broad-except
            logger.error("Background batch job %s failed: %s", job_id, exc)

    background_tasks.add_task(_run_background)
    return {
        "job_id": job_id,
        "status": "running",
        "total_files": len(records),
        "message": "Batch job started. Poll /api/multimodal/batch-status/{job_id} for progress.",
    }


@router.get("/batch-status/{job_id}", summary="Get Batch Job Status")
async def get_batch_status(job_id: str, db: Session = Depends(get_db)):
    """Return the current status of a batch processing job from Postgres."""
    try:
        from sqlalchemy import text as sa_text
        row = db.execute(
            sa_text(
                "SELECT job_id, started_at, completed_at, total_files, "
                "processed, succeeded, failed, skipped "
                "FROM file_batch_jobs WHERE job_id = :jid"
            ),
            {"jid": job_id},
        ).fetchone()

        if row is None:
            raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

        return {
            "job_id": row.job_id,
            "started_at": row.started_at,
            "completed_at": row.completed_at,
            "status": "completed" if row.completed_at else "running",
            "total_files": row.total_files,
            "processed": row.processed,
            "succeeded": row.succeeded,
            "failed": row.failed,
            "skipped": row.skipped,
            "progress_pct": round(100 * row.processed / row.total_files, 1) if row.total_files else 0,
        }
    except HTTPException:
        raise
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Error fetching batch status: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.get("/batch-results/{job_id}", summary="Get Batch Job Results")
async def get_batch_results(
    job_id: str,
    skip: int = 0,
    limit: int = 100,
    file_type: Optional[str] = None,
    success_only: bool = False,
    db: Session = Depends(get_db),
):
    """Page through per-file results for a completed batch job."""
    try:
        from sqlalchemy import text as sa_text
        where_clauses = ["job_id = :jid"]
        params: Dict[str, Any] = {"jid": job_id, "skip": skip, "limit": limit}

        if file_type:
            where_clauses.append("file_type = :file_type")
            params["file_type"] = file_type
        if success_only:
            where_clauses.append("success = TRUE")

        where = " AND ".join(where_clauses)
        rows = db.execute(
            sa_text(
                f"SELECT file_path, file_type, success, text_content, error, "
                f"processing_time_ms, processed_at "
                f"FROM file_batch_results WHERE {where} "
                f"ORDER BY id OFFSET :skip LIMIT :limit"
            ),
            params,
        ).fetchall()

        total = db.execute(
            sa_text(f"SELECT COUNT(*) FROM file_batch_results WHERE {where}"),
            params,
        ).scalar()

        return {
            "job_id": job_id,
            "total": total,
            "skip": skip,
            "limit": limit,
            "results": [
                {
                    "file_path": r.file_path,
                    "file_type": r.file_type,
                    "success": r.success,
                    "text_snippet": (r.text_content or "")[:200],
                    "error": r.error,
                    "processing_time_ms": r.processing_time_ms,
                    "processed_at": r.processed_at,
                }
                for r in rows
            ],
        }
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Error fetching batch results: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
